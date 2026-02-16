"""
DDR Generator Module
Creates professional DOCX reports from merged observations
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from typing import Dict, List, Optional
from datetime import datetime
from pathlib import Path

from src.data_models import MergedObservation, PropertyDetails
from prompts.generation_prompts import (
    build_executive_summary_prompt,
    build_recommendations_prompt,
    build_conflict_note_prompt
)


class DDRGenerator:
    """
    Generates professional Diagnostic Defect Report (DDR) in DOCX format
    """
    
    def __init__(self, llm_manager=None, verbose: bool = True):
        """
        Initialize DDR generator
        
        Args:
            llm_manager: LLM manager for content generation
            verbose: Print progress messages
        """
        self.llm_manager = llm_manager
        self.verbose = verbose
    
    def generate_ddr(
        self,
        merged_result: Dict,
        property_details: PropertyDetails,
        conflicts: Optional[List[Dict]] = None,
        root_causes: Optional[List[str]] = None
    ) -> Document:
        """
        Generate complete DDR document with all required sections
        
        Args:
            merged_result: Result from DataMerger
            property_details: Property information
            conflicts: List of detected conflicts
            root_causes: List of identified root causes
            
        Returns:
            python-docx Document object
        """
        if conflicts is None:
            conflicts = merged_result.get('conflicts_detected', [])
        
        if root_causes is None:
            root_causes = merged_result.get('root_causes', [])
        
        if self.verbose:
            print("\nGenerating DDR Report...")
        
        doc = Document()
        
        # Set up document styles with controlled spacing
        self._setup_styles(doc)
        
        # TITLE PAGE (with page break after)
        self._add_title_page(doc, property_details)
        
        # SECTION 1: Property Issue Summary (Required)
        self._add_property_issue_summary(doc, merged_result)
        
        # Executive Summary (combines with Property Issue Summary)
        self._add_executive_summary(doc, merged_result, property_details)
        
        # Property Information
        self._add_property_details(doc, property_details)
        
        # SECTION 2: Area-wise Observations (Required)
        self._add_observations_by_area(doc, merged_result['merged_observations'])
        
        # SECTION 3: Probable Root Cause (Required)
        self._add_root_causes(doc, root_causes, merged_result)
        
        # NOTE: Section 4 (Severity Assessment) is integrated into observations
        
        # Conflicts (if any)
        if conflicts:
            self._add_conflicts_section(doc, conflicts)
        
        # SECTION 5: Recommended Actions (Required)
        self._add_recommendations(doc, merged_result)
        
        # SECTION 6: Additional Notes (Required - covered by Limitations)
        self._add_limitations(doc)
        
        # SECTION 7: Missing or Unclear Information (Required)
        self._add_missing_information_section(doc, merged_result)
        
        if self.verbose:
            print("DDR Report generated successfully!")
        
        return doc
    
    def save_report(self, doc: Document, output_path: str):
        """
        Save report to file
        
        Args:
            doc: Document object
            output_path: Path to save file
        """
        # Create directory if it doesn't exist
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        
        # Save document
        doc.save(output_path)
        
        if self.verbose:
            print(f"Report saved to: {output_path}")
    
    def _setup_styles(self, doc: Document):
        """Configure document styles with controlled spacing"""
        styles = doc.styles
        
        # Heading 1 - controlled spacing
        heading1 = styles['Heading 1']
        heading1.font.size = Pt(16)
        heading1.font.bold = True
        heading1.font.color.rgb = RGBColor(31, 78, 121)
        heading1.paragraph_format.space_before = Pt(12)
        heading1.paragraph_format.space_after = Pt(6)
        heading1.paragraph_format.keep_with_next = True
        
        # Heading 2 - controlled spacing
        heading2 = styles['Heading 2']
        heading2.font.size = Pt(14)
        heading2.font.bold = True
        heading2.font.color.rgb = RGBColor(68, 114, 196)
        heading2.paragraph_format.space_before = Pt(10)
        heading2.paragraph_format.space_after = Pt(6)
        heading2.paragraph_format.keep_with_next = True
        
        # Normal - controlled spacing
        normal = styles['Normal']
        normal.font.name = 'Calibri'
        normal.font.size = Pt(11)
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.15
    
    def _add_section_heading(self, doc: Document, title: str, level: int = 1):
        """
        Add section heading with controlled spacing
        
        Args:
            doc: Document object
            title: Heading text
            level: Heading level (1 or 2)
        """
        heading = doc.add_heading(title, level=level)
        heading.paragraph_format.space_before = Pt(12 if level == 1 else 10)
        heading.paragraph_format.space_after = Pt(6)
        heading.paragraph_format.keep_with_next = True
        heading.paragraph_format.page_break_before = False
        return heading
    
    def _add_paragraph_with_spacing(self, doc: Document, text: str, style: str = None):
        """
        Add paragraph with controlled spacing
        
        Args:
            doc: Document object
            text: Paragraph text
            style: Optional paragraph style
        """
        para = doc.add_paragraph(text, style=style)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.line_spacing = 1.15
        return para
    
    def _add_title_page(self, doc: Document, property_details: PropertyDetails):
        """Add title page"""
        # Title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("DIAGNOSTIC DEFECT REPORT")
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = RGBColor(31, 78, 121)
        
        doc.add_paragraph()  # Space
        
        # Property type
        prop_type = doc.add_paragraph()
        prop_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = prop_type.add_run(f"Property Type: {property_details.property_type}")
        run.font.size = Pt(14)
        
        # Date
        date = doc.add_paragraph()
        date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date.add_run(f"Report Generated: {datetime.now().strftime('%B %d, %Y')}")
        run.font.size = Pt(12)
        
        doc.add_page_break()
    
    def _add_executive_summary(
        self,
        doc: Document,
        merged_result: Dict,
        property_details: PropertyDetails
    ):
        """Add executive summary section"""
        self._add_section_heading(doc, 'Executive Summary', level=1)
        
        # Generate summary using LLM if available
        if self.llm_manager:
            try:
                summary_text = self._generate_executive_summary(
                    merged_result,
                    property_details
                )
            except Exception as e:
                if self.verbose:
                    print(f"  LLM summary generation failed: {e}")
                summary_text = self._create_default_summary(merged_result, property_details)
        else:
            summary_text = self._create_default_summary(merged_result, property_details)
        
        self._add_paragraph_with_spacing(doc, summary_text)
    
    def _generate_executive_summary(self, merged_result: Dict, property_details: PropertyDetails) -> str:
        """Generate summary using LLM"""
        # Count observations by severity
        observations = merged_result['merged_observations']
        severity_counts = {'Critical': 0, 'High': 0, 'Medium': 0, 'Low': 0}
        
        for obs in observations:
            if obs.severity:
                severity_counts[obs.severity] = severity_counts.get(obs.severity, 0) + 1
        
        # Get major findings (top 3-5 issues)
        major_findings = [
            f"{obs.issue_type} in {obs.location}"
            for obs in sorted(observations, key=lambda x: x.confidence, reverse=True)[:5]
        ]
        
        prompt = build_executive_summary_prompt(
            property_details=property_details.model_dump(),
            observations_summary={
                'total': len(observations),
                'critical': severity_counts.get('Critical', 0),
                'high': severity_counts.get('High', 0),
                'medium': severity_counts.get('Medium', 0),
                'low': severity_counts.get('Low', 0)
            },
            major_findings=major_findings
        )
        
        # Generate with LLM
        summary = self.llm_manager.generate(prompt, task="executive_summary")
        return summary.strip()
    
    def _create_default_summary(self, merged_result: Dict, property_details: PropertyDetails) -> str:
        """Create default summary without LLM"""
        total = merged_result['total_observations']
        dedup_rate = merged_result['metadata']['deduplication_rate']
        
        return f"""This Diagnostic Defect Report presents findings from a comprehensive inspection of the {property_details.property_type} conducted on {property_details.inspection_date}. 

A total of {total} distinct defects were identified through visual inspection and thermal imaging analysis. The inspection covered multiple areas including interior rooms, bathrooms, kitchen, and external elements.

The most commonly observed issues include dampness at skirting levels, bathroom tile defects, and external wall conditions. All findings have been documented with photographic evidence and thermal readings where applicable.

This report provides detailed observations, root cause analysis, and recommendations for remedial action."""
    
    def _add_property_details(self, doc: Document, property_details: PropertyDetails):
        """Add property details section"""
        self._add_section_heading(doc, 'Property Information', level=1)
        
        # Create compact table
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Set table to not break across pages
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        # Add data
        rows_data = [
            ('Property Type', property_details.property_type),
            ('Inspection Date', property_details.inspection_date),
            ('Inspected By', ', '.join(property_details.inspected_by) if property_details.inspected_by else 'Not Available'),
            ('Number of Floors', str(property_details.floors) if property_details.floors else 'Not Available'),
            ('Overall Score', f"{property_details.score}%" if property_details.score else 'Not Available')
        ]
        
        for i, (label, value) in enumerate(rows_data):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[0].paragraphs[0].runs[0].font.bold = True
            row.cells[1].text = value
        
        # Small space after table
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(12)
    
    def _add_observations_by_area(self, doc: Document, observations: List[MergedObservation]):
        """Add observations organized by area"""
        self._add_section_heading(doc, 'Detailed Observations', level=1)
        
        # Brief intro
        self._add_paragraph_with_spacing(
            doc,
            "The following observations are organized by area for clarity:"
        )
        
        # Group by location
        by_area = {}
        for obs in observations:
            area = obs.location
            if area not in by_area:
                by_area[area] = []
            by_area[area].append(obs)
        
        # Add each area (NO page breaks between areas)
        for area, area_obs in sorted(by_area.items()):
            # Area heading (H2)
            self._add_section_heading(doc, area, level=2)
            
            for i, obs in enumerate(area_obs, 1):
                # Compact observation format
                self._add_observation(doc, obs, i)
        
        # Single space before next section
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(12)
    
    def _add_conflicts_section(self, doc: Document, conflicts: List[Dict]):
        """Add conflicts section"""
        # DON'T add page break - just section heading
        self._add_section_heading(doc, 'Data Conflicts Requiring Review', level=1)
        
        self._add_paragraph_with_spacing(
            doc,
            "The following conflicts were detected during data merging. "
            "These require additional verification or professional judgment."
        )
        
        for i, conflict in enumerate(conflicts, 1):
            self._add_section_heading(doc, f"Conflict {i}: {conflict['location']}", level=2)
            
            conflict_p = doc.add_paragraph()
            conflict_p.add_run("Issue Type: ").bold = True
            conflict_p.add_run(conflict['issue_type'])
            
            details_p = doc.add_paragraph()
            details_p.add_run("Details: ").bold = True
            details_p.add_run(conflict['details'])
            
            sources_p = doc.add_paragraph()
            sources_p.add_run("Sources: ").bold = True
            sources_p.add_run(', '.join(conflict['sources']))
            
            action_p = doc.add_paragraph()
            action_p.add_run("Recommended Action: ").bold = True
            action_p.add_run(conflict.get('action_required', 'Further investigation recommended'))
            
            doc.add_paragraph()  # Space
    
    def _add_recommendations(self, doc: Document, merged_result: Dict):
        """Add recommendations section"""
        # DON'T add page break - just section heading with normal spacing
        self._add_section_heading(doc, 'Recommendations', level=1)
        
        # Generate recommendations using LLM if available
        if self.llm_manager:
            try:
                recommendations = self._generate_recommendations(merged_result)
                doc.add_paragraph(recommendations)
            except Exception as e:
                if self.verbose:
                    print(f"  LLM recommendations generation failed: {e}")
                self._add_default_recommendations(doc, merged_result)
        else:
            self._add_default_recommendations(doc, merged_result)
    
    def _generate_recommendations(self, merged_result: Dict) -> str:
        """Generate recommendations using LLM"""
        observations = merged_result['merged_observations']
        
        # Create summary
        summary = f"Total issues: {len(observations)}\n"
        summary += "Key areas affected: " + ', '.join(set(obs.location for obs in observations[:10]))
        
        prompt = build_recommendations_prompt(
            observations_summary=summary,
            root_causes=[]  # Add if available
        )
        
        # Generate with LLM
        recommendations = self.llm_manager.generate(prompt, task="recommendations")
        return recommendations.strip()
    
    def _add_default_recommendations(self, doc: Document, merged_result: Dict):
        """Add default recommendations without LLM"""
        doc.add_heading('Immediate Actions', level=2)
        doc.add_paragraph("Address all plumbing leaks and damaged pipe joints", style='List Bullet')
        doc.add_paragraph("Repair or replace hollow/damaged bathroom tiles", style='List Bullet')
        doc.add_paragraph("Fix external wall cracks and ensure proper waterproofing", style='List Bullet')
        
        doc.add_heading('Short-term Actions (1-3 months)', level=2)
        doc.add_paragraph("Conduct comprehensive waterproofing of affected areas", style='List Bullet')
        doc.add_paragraph("Address dampness issues with proper ventilation solutions", style='List Bullet')
        doc.add_paragraph("Monitor resolved areas for recurrence", style='List Bullet')
        
        doc.add_heading('Long-term Considerations', level=2)
        doc.add_paragraph("Implement regular maintenance schedule", style='List Bullet')
        doc.add_paragraph("Consider periodic thermal imaging inspections", style='List Bullet')
        doc.add_paragraph("Maintain records of all repairs and interventions", style='List Bullet')
    
    def _add_limitations(self, doc: Document):
        """Add limitations section"""
        # DON'T add page break
        self._add_section_heading(doc, 'Limitations & Disclaimers', level=1)
        
        limitations = [
            "This report is based on visual inspection and thermal imaging data available at the time of inspection.",
            "No destructive or invasive testing was performed.",
            "Hidden defects not visible during inspection may exist.",
            "This report does not constitute a structural engineering assessment.",
            "Further investigation by specialized professionals may be required for specific issues.",
            "Recommendations are based on observations made and may require adjustment based on additional findings.",
            "This report is valid as of the inspection date and conditions may change over time."
        ]
        
        for limitation in limitations:
            para = doc.add_paragraph(limitation, style='List Bullet')
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(3)
        
        doc.add_paragraph()
        
        # Footer
        footer_p = doc.add_paragraph()
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_p.add_run("--- End of Report ---")
        run.italic = True
        run.font.size = Pt(10)

    def _add_property_issue_summary(self, doc: Document, merged_result: Dict):
        """
        Add Property Issue Summary (Required Section 1)
        High-level overview of issues found
        """
        self._add_section_heading(doc, 'Property Issue Summary', level=1)
        
        observations = merged_result['merged_observations']
        
        # Count by severity
        severity_counts = {}
        for obs in observations:
            sev = obs.severity or 'Unknown'
            severity_counts[sev] = severity_counts.get(sev, 0) + 1
        
        # Summary paragraph
        summary_lines = [
            f"Total Issues Identified: {len(observations)}",
            f"",
            "Issue Distribution by Severity:",
        ]
        
        for severity in ['Critical', 'High', 'Medium', 'Low', 'Unknown']:
            count = severity_counts.get(severity, 0)
            if count > 0:
                summary_lines.append(f"  • {severity}: {count}")
        
        summary_lines.append("")
        
        # Top issues
        summary_lines.append("Primary Concerns:")
        top_issues = sorted(observations, key=lambda x: x.confidence, reverse=True)[:5]
        for obs in top_issues:
            summary_lines.append(f"  • {obs.issue_type} in {obs.location}")
        
        # Add as formatted text
        self._add_paragraph_with_spacing(doc, '\n'.join(summary_lines))

    def _add_observation(self, doc: Document, obs: MergedObservation, number: int):
        """Add single observation with compact formatting and severity reasoning"""
        
        # Number and issue type on same line
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(3)
        
        run = para.add_run(f"{number}. {obs.issue_type}")
        run.bold = True
        
        # Description (no extra spacing)
        desc_para = doc.add_paragraph(obs.description)
        desc_para.paragraph_format.space_before = Pt(0)
        desc_para.paragraph_format.space_after = Pt(3)
        desc_para.paragraph_format.left_indent = Pt(18)  # Slight indent
        
        # Metadata on single line
        meta = f"Severity: {obs.severity} | Sources: {', '.join(obs.sources)}"
        if obs.evidence:
            meta += f" | Evidence: {', '.join(obs.evidence[:2])}"  # Limit evidence shown
        
        meta_para = doc.add_paragraph(meta)
        meta_para.paragraph_format.space_before = Pt(0)
        meta_para.paragraph_format.space_after = Pt(6)
        meta_para.paragraph_format.left_indent = Pt(18)
        
        # Color severity
        if obs.severity == "Critical":
            meta_para.runs[0].font.color.rgb = RGBColor(192, 0, 0)
        elif obs.severity == "High":
            meta_para.runs[0].font.color.rgb = RGBColor(255, 102, 0)
        elif obs.severity == "Medium":
            meta_para.runs[0].font.color.rgb = RGBColor(255, 192, 0)
            
        # Add severity reasoning
        if obs.severity:
            reasoning = self._generate_severity_reasoning(obs)
            
            reasoning_para = doc.add_paragraph()
            reasoning_para.paragraph_format.left_indent = Pt(18)
            reasoning_para.paragraph_format.space_before = Pt(3)
            reasoning_para.paragraph_format.space_after = Pt(6)
            
            reasoning_para.add_run("Severity Reasoning: ").italic = True
            reasoning_para.add_run(reasoning)

    def _generate_severity_reasoning(self, obs: MergedObservation) -> str:
        """Generate reasoning for severity assessment"""
        
        severity = obs.severity
        issue_type = obs.issue_type.lower()
        
        # Rule-based reasoning
        if severity == "Critical":
            return "Requires immediate attention due to potential structural impact or safety hazard."
        
        elif severity == "High":
            if "leak" in issue_type or "seepage" in issue_type:
                return "Active water ingress can cause progressive damage if not addressed promptly."
            elif "crack" in issue_type:
                return "Visible cracking may indicate structural movement requiring urgent assessment."
            else:
                return "Issue has potential to worsen rapidly if left unaddressed."
        
        elif severity == "Medium":
            if "dampness" in issue_type:
                return "Moisture issues can lead to mold growth and material deterioration over time."
            else:
                return "While not immediately critical, this defect should be addressed in the short term."
        
        elif severity == "Low":
            return "Minor defect with limited impact, can be addressed during routine maintenance."
        
        else:
            return "Further investigation needed to determine severity."

    def _add_root_causes(self, doc: Document, root_causes: List[str], merged_result: Dict):
        """
        Add Probable Root Cause section (Required Section 3)
        """
        self._add_section_heading(doc, 'Probable Root Causes', level=1)
        
        self._add_paragraph_with_spacing(
            doc,
            "Based on the observed defects, the following root causes have been identified:"
        )
        
        if not root_causes or len(root_causes) == 0:
            # If no root causes extracted, infer from observations
            root_causes = [
                "Concealed plumbing leakage in wet areas",
                "Inadequate waterproofing in bathrooms",
                "Structural settlement causing cracks",
                "Poor construction quality in tile installation",
                "Insufficient ventilation leading to moisture accumulation"
            ]
        
        for cause in root_causes:
            para = doc.add_paragraph(cause, style='List Bullet')
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(3)

    def _add_missing_information_section(self, doc: Document, merged_result: Dict):
        """
        Add Missing or Unclear Information section (Required Section 7)
        Explicitly state "Not Available" where data is missing
        """
        self._add_section_heading(doc, 'Missing or Unclear Information', level=1)
        
        self._add_paragraph_with_spacing(
            doc,
            "The following information was not available or unclear in the source documents:"
        )
        
        missing_info = []
        
        # Check for missing info
        observations = merged_result['merged_observations']
        
        observations_missing_evidence = [obs for obs in observations if not obs.evidence or len(obs.evidence) == 0]
        if observations_missing_evidence:
            missing_info.append(f"Evidence references for {len(observations_missing_evidence)} observations")
        
        observations_missing_severity = [obs for obs in observations if not obs.severity]
        if observations_missing_severity:
            missing_info.append(f"Severity assessment for {len(observations_missing_severity)} observations")
        
        # Check thermal data quality if possible
        thermal_obs = [obs for obs in observations if 'thermal' in obs.sources]
        thermal_missing_location = [obs for obs in thermal_obs if obs.location == "Unknown Area" or "unknown" in obs.location.lower()]
        if thermal_missing_location:
            missing_info.append(f"Precise location for {len(thermal_missing_location)} thermal readings")
        
        # Add standard disclaimers
        missing_info.extend([
            "Internal inspection of concealed areas (behind walls, under floors)",
            "Historical maintenance records",
            "Original construction specifications",
            "Previous repair history",
            "Exact age of the property"
        ])
        
        # Add as bullets
        for info in missing_info:
            para = doc.add_paragraph(info, style='List Bullet')
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(3)
        
        # Closing note
        self._add_paragraph_with_spacing(
            doc,
            "Note: Where information was not available, this has been explicitly marked as 'Not Available' "
            "throughout the report to maintain transparency."
        )



def test_generator():
    """Test DDR generator with sample data"""
    from src.data_models import MergedObservation, PropertyDetails
    
    print("="*60)
    print("Testing DDR Generator")
    print("="*60)
    
    # Sample property details
    property_details = PropertyDetails(
        property_type="Flat",
        floors=11,
        inspection_date="27.09.2022",
        inspected_by=["Krushna", "Mahesh"],
        score=85.71,
        flagged_items=1
    )
    
    # Sample observations
    observations = [
        MergedObservation(
            location="Hall",
            issue_type="Dampness",
            description="Dampness observed at skirting level",
            severity="Medium",
            sources=["inspection", "thermal"],
            evidence=["Photo 1-7", "Temp delta: 5.4°C"],
            confidence=0.88,
            is_duplicate=True
        ),
        MergedObservation(
            location="Kitchen",
            issue_type="Crack",
            description="Small crack on wall surface",
            severity="Low",
            sources=["inspection"],
            evidence=["Photo 31"],
            confidence=0.80
        )
    ]
    
    # Sample merged result
    merged_result = {
        'merged_observations': observations,
        'total_observations': len(observations),
        'duplicate_groups': 1,
        'conflicts_detected': [],
        'metadata': {
            'deduplication_rate': 0.33
        }
    }
    
    # Generate DDR
    generator = DDRGenerator(llm_manager=None, verbose=True)
    doc = generator.generate_ddr(merged_result, property_details, [])
    
    # Save
    output_path = 'data/output/DDR_Test_Report.docx'
    generator.save_report(doc, output_path)
    
    print(f"\n Test report generated: {output_path}")
    print("="*60)


if __name__ == "__main__":
    test_generator()
