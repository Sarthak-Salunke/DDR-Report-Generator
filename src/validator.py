"""
Quality Validator Module
Validates the quality and completeness of generated DDR reports
"""

from pathlib import Path
from typing import Dict, List, Tuple, Optional
import json
from docx import Document

from src.data_models import MergedObservation


class QualityValidator:
    """
    Validates DDR report quality and completeness
    """
    
    def __init__(self, verbose: bool = True):
        self.verbose = verbose
    
    def validate_report(
        self,
        docx_path: str,
        merged_result: Dict,
        source_texts: Optional[List[str]] = None
    ) -> Dict:
        """
        Comprehensive report validation
        
        Args:
            docx_path: Path to generated DOCX file
            merged_result: Merged observations data
            source_texts: Original source document texts (optional)
            
        Returns:
            Validation report dictionary
        """
        if self.verbose:
            print("\n✓ Running quality validation...")
        
        validation_results = {
            'passed': True,
            'errors': [],
            'warnings': [],
            'metrics': {}
        }
        
        # Check 1: File exists and is valid
        file_check = self._validate_file(docx_path)
        validation_results['errors'].extend(file_check['errors'])
        validation_results['warnings'].extend(file_check['warnings'])
        validation_results['metrics']['file_size_kb'] = file_check['size_kb']
        
        # Check 2: Document structure
        if Path(docx_path).exists():
            doc = Document(docx_path)
            structure_check = self._validate_structure(doc)
            validation_results['errors'].extend(structure_check['errors'])
            validation_results['warnings'].extend(structure_check['warnings'])
            validation_results['metrics'].update(structure_check['metrics'])
        
        # Check 3: Content completeness
        completeness_check = self._validate_completeness(
            merged_result,
            docx_path if Path(docx_path).exists() else None
        )
        validation_results['errors'].extend(completeness_check['errors'])
        validation_results['warnings'].extend(completeness_check['warnings'])
        
        # Check 4: Data consistency
        consistency_check = self._validate_consistency(merged_result)
        validation_results['errors'].extend(consistency_check['errors'])
        validation_results['warnings'].extend(consistency_check['warnings'])
        
        # Determine if passed
        validation_results['passed'] = len(validation_results['errors']) == 0
        
        if self.verbose:
            self._print_validation_report(validation_results)
        
        return validation_results
    
    def _validate_file(self, docx_path: str) -> Dict:
        """Validate file existence and basic properties"""
        errors = []
        warnings = []
        
        if not Path(docx_path).exists():
            errors.append(f"Output file does not exist: {docx_path}")
            return {'errors': errors, 'warnings': warnings, 'size_kb': 0}
        
        file_size = Path(docx_path).stat().st_size
        size_kb = file_size / 1024
        
        if size_kb < 10:
            errors.append(f"File size too small ({size_kb:.1f} KB), likely incomplete")
        elif size_kb < 20:
            warnings.append(f"File size small ({size_kb:.1f} KB), may be missing content")
        
        # Try to open as DOCX
        try:
            doc = Document(docx_path)
            if len(doc.paragraphs) < 10:
                errors.append("Document has too few paragraphs")
        except Exception as e:
            errors.append(f"Cannot open as valid DOCX: {e}")
        
        return {'errors': errors, 'warnings': warnings, 'size_kb': size_kb}
    
    def _validate_structure(self, doc: Document) -> Dict:
        """Validate document structure"""
        errors = []
        warnings = []
        metrics = {}
        
        # Extract all text
        all_text = '\n'.join([p.text for p in doc.paragraphs])
        
        # Required sections
        required_sections = [
            'Executive Summary',
            'Property Information',
            'Detailed Observations',
            'Recommendations',
            'Limitations'
        ]
        
        for section in required_sections:
            if section not in all_text:
                errors.append(f"Missing required section: {section}")
        
        # Count elements
        metrics['total_paragraphs'] = len(doc.paragraphs)
        metrics['total_tables'] = len(doc.tables)
        metrics['total_characters'] = len(all_text)
        
        # Check for minimum content
        if metrics['total_paragraphs'] < 20:
            warnings.append(f"Few paragraphs ({metrics['total_paragraphs']}), expected 20+")
        
        if metrics['total_tables'] < 1:
            warnings.append("No tables found, property details may be missing")
        
        return {'errors': errors, 'warnings': warnings, 'metrics': metrics}
    
    def _validate_completeness(self, merged_result: Dict, docx_path: Optional[str] = None) -> Dict:
        """Validate data completeness"""
        errors = []
        warnings = []
        
        # Check observations count
        total_obs = merged_result['total_observations']
        if total_obs < 5:
            errors.append(f"Too few observations ({total_obs}), expected at least 5")
        elif total_obs < 8:
            warnings.append(f"Low observation count ({total_obs}), expected 8+")
        
        # Check for empty descriptions
        empty_descriptions = 0
        for obs in merged_result['merged_observations']:
            if not obs.description or len(obs.description) < 10:
                empty_descriptions += 1
        
        if empty_descriptions > 0:
            warnings.append(f"{empty_descriptions} observations have very short descriptions")
        
        # Check for source diversity
        sources_used = set()
        for obs in merged_result['merged_observations']:
            sources_used.update(obs.sources)
        
        if len(sources_used) < 2:
            warnings.append("Only one source type used, expected both inspection and thermal")
        
        return {'errors': errors, 'warnings': warnings}
    
    def _validate_consistency(self, merged_result: Dict) -> Dict:
        """Validate data consistency"""
        errors = []
        warnings = []
        
        # Check for duplicate locations with inconsistent data
        by_location = {}
        for obs in merged_result['merged_observations']:
            loc = obs.location
            if loc not in by_location:
                by_location[loc] = []
            by_location[loc].append(obs)
        
        # Check severity consistency within locations
        for location, obs_list in by_location.items():
            severities = [obs.severity for obs in obs_list if obs.severity]
            if len(set(severities)) > 2:
                warnings.append(
                    f"Location '{location}' has inconsistent severities: {set(severities)}"
                )
        
        # Check confidence scores
        low_confidence = [
            obs for obs in merged_result['merged_observations']
            if obs.confidence < 0.6
        ]
        
        if low_confidence:
            warnings.append(
                f"{len(low_confidence)} observations have low confidence (<0.6)"
            )
        
        return {'errors': errors, 'warnings': warnings}
    
    def _print_validation_report(self, results: Dict):
        """Print formatted validation report"""
        print("\n" + "="*70)
        print("📋 VALIDATION REPORT")
        print("="*70)
        
        # Overall status
        if results['passed']:
            print("✅ VALIDATION PASSED")
        else:
            print("❌ VALIDATION FAILED")
        
        print()
        
        # Metrics
        if results['metrics']:
            print("METRICS:")
            for key, value in results['metrics'].items():
                print(f"  • {key}: {value}")
            print()
        
        # Errors
        if results['errors']:
            print(f"ERRORS ({len(results['errors'])}):")
            for error in results['errors']:
                print(f"  ❌ {error}")
            print()
        
        # Warnings
        if results['warnings']:
            print(f"WARNINGS ({len(results['warnings'])}):")
            for warning in results['warnings']:
                print(f"  ⚠️  {warning}")
            print()
        
        if not results['errors'] and not results['warnings']:
            print("✅ No issues found!")
            print()
        
        print("="*70)


def test_validator():
    """Test the validator"""
    from src.data_models import MergedObservation
    
    print("="*70)
    print("Testing Quality Validator")
    print("="*70)
    
    # Sample merged result
    sample_obs = [
        MergedObservation(
            location="Hall",
            issue_type="Dampness",
            description="Dampness at skirting level",
            sources=["inspection", "thermal"],
            evidence=["Photo1"],
            confidence=0.9
        ),
        MergedObservation(
            location="Kitchen",
            issue_type="Crack",
            description="Wall crack detected",
            sources=["inspection"],
            evidence=[],
            confidence=0.8
        )
    ]
    
    merged_result = {
        'merged_observations': sample_obs,
        'total_observations': len(sample_obs),
        'duplicate_groups': 0,
        'conflicts_detected': [],
        'metadata': {}
    }
    
    validator = QualityValidator(verbose=True)
    
    # Test with actual file if exists
    test_file = 'data/output/DDR_Test_Report.docx'
    if Path(test_file).exists():
        print(f"\n✓ Found test report: {test_file}")
        result = validator.validate_report(
            test_file,
            merged_result,
            ["source text 1", "source text 2"]
        )
    else:
        print(f"\n⚠️  Test report not found: {test_file}")
        print("Testing with non-existent file...")
        result = validator.validate_report(
            "nonexistent.docx",
            merged_result,
            []
        )
    
    print("\n✅ Validator test completed!")
    return result


if __name__ == "__main__":
    test_validator()
